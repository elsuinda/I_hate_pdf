import os
from docx import Document
from openpyxl import load_workbook
from fpdf import FPDF

class FileConverter:
    @staticmethod
    def convert(filepath, target_format):
        if not os.path.exists(filepath):
            raise FileNotFoundError(f"El archivo '{filepath}' no existe.")

        filename, ext = os.path.splitext(filepath)
        ext = ext.lower()  # Asegurarse de que la extensión sea minúscula
        target_format = target_format.lower()

        # Mapear formatos a extensiones correctas
        extension_map = {
            "word": "docx",
            "excel": "xlsx",
            "pdf": "pdf"
        }

        if target_format not in extension_map:
            raise ValueError(f"Formato de destino '{target_format}' no soportado.")

        output_path = f"{filename}.{extension_map[target_format]}"

        print(f"Converting {filepath} to {target_format}...")

        # Validar formatos soportados
        supported_conversions = {
            ".pdf": ["word", "excel"],
            ".docx": ["pdf"],
            ".xlsx": ["pdf"]
        }

        if ext not in supported_conversions or target_format not in supported_conversions[ext]:
            raise ValueError(f"Conversión de {ext} a {target_format} no soportada.")

        # Lógica de conversión real
        if ext == ".docx" and target_format == "pdf":
            doc = Document(filepath)
            pdf = FPDF()
            pdf.set_auto_page_break(auto=True, margin=15)
            pdf.add_page()
            pdf.set_font("Arial", size=12)
            for paragraph in doc.paragraphs:
                pdf.multi_cell(0, 10, paragraph.text)
            pdf.output(output_path)
        elif ext == ".xlsx" and target_format == "pdf":
            wb = load_workbook(filepath)
            ws = wb.active
            pdf = FPDF()
            pdf.set_auto_page_break(auto=True, margin=15)
            pdf.add_page()
            pdf.set_font("Arial", size=12)
            for row in ws.iter_rows(values_only=True):
                row_text = " | ".join([str(cell) if cell is not None else "" for cell in row])
                pdf.multi_cell(0, 10, row_text)
            pdf.output(output_path)
        elif ext == ".pdf" and target_format == "word":
            doc = Document()
            doc.add_heading("Archivo convertido", level=1)
            try:
                reader = PdfReader(filepath)
                for page in reader.pages:
                    doc.add_paragraph(page.extract_text())
            except Exception:
                doc.add_paragraph("No se pudo leer el contenido del archivo original.")
            doc.save(output_path)

        print("Conversion complete.")
        return output_path
